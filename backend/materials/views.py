from rest_framework import viewsets, status
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.decorators import action, permission_classes
from classrooms.models import Classroom
from .models import Document
from .serializers import DocumentSerializer
import pdfplumber
import docx
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
import tempfile
import os
from .models import Document
from django.core.exceptions import ObjectDoesNotExist
from django.conf import settings


class DocumentViewSet(viewsets.ModelViewSet):
    serializer_class = DocumentSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        queryset = Document.objects.filter(classroom__creator=self.request.user)

        # Filter by classroom if provided
        classroom_id = self.request.query_params.get('classroom_id')
        if classroom_id:
            queryset = queryset.filter(classroom_id=classroom_id)

        # Filter by tag names if provided
        tag_names = self.request.query_params.get('tag_names')
        if tag_names:
            # Split the comma-separated tag names
            tag_names_list = tag_names.split(',')
            # Filter documents that have ALL the specified tags
            for tag_name in tag_names_list:
                queryset = queryset.filter(tags__name=tag_name)

        return queryset

    def create(self, request, *args, **kwargs):
        classroom_id = request.data.get('classroom')
        if classroom_id:
            try:
                classroom = Classroom.objects.get(id=classroom_id)

                if classroom.creator != request.user:
                    return Response(
                        {"error": "You don't have permission to add documents to this classroom."},
                        status=status.HTTP_400_BAD_REQUEST
                    )

                if classroom.documents.count() >= settings.MAX_FILES_PER_CLASSROOM:
                    return Response(
                        {"error": f"This classroom already has the maximum number of files ({settings.MAX_FILES_PER_CLASSROOM})."},
                        status=status.HTTP_400_BAD_REQUEST
                    )
            except ObjectDoesNotExist:
                return Response(
                    {"error": "Classroom does not exist."},
                    status=status.HTTP_400_BAD_REQUEST
                )
        return super().create(request, *args, **kwargs)

    def update_tags(self, request, pk=None):
        document = self.get_object()
        tag_ids = request.data.get('tag_ids', [])

        try:
            document.tags.set(tag_ids)
            serializer = self.get_serializer(document)
            return Response(serializer.data)
        except Exception as e:
            return Response(
                {'error': f"Failed to update tags: {str(e)}"},  # Added detailed error message
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=False, methods=['get'], url_path='user_materials')
    def get_all_user_materials(self, request):
        try:
            materials = Document.objects.filter(
                classroom__creator=request.user
            ).select_related('classroom').prefetch_related('tags')
            serializer = DocumentSerializer(materials, many=True)
            return Response(serializer.data)
        except Exception as e:
            return Response(
                {'error': f"Failed to fetch user materials: {str(e)}"},  # Added detailed error message
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=False, methods=['post'], url_path='extract-text')
    def extract_text_from_uploaded_file(self, request):
        """Extract text from an uploaded DOCX file"""
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=400)
        
        file = request.FILES['file']
        
        # Check if it's a DOCX
        if not file.name.lower().endswith('.docx'):
            return Response({'error': 'Only DOCX files are supported'}, status=400)
        
        try:
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                for chunk in file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name
            
            # Extract text using python-docx
            doc = docx.Document(temp_file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"
            
            # Clean up the temporary file
            os.unlink(temp_file_path)
            
            return Response({'text': text})
        
        except Exception as e:
            return Response({'error': str(e)}, status=500)

    @action(detail=True, methods=['get'], url_path='extract-text')
    def extract_text_from_material(self, request, pk=None):
        """Extract text from a material that's already in the database"""
        try:
            # Get the material
            material = self.get_object()
            
            # Check if it's a DOCX
            if not material.file.name.lower().endswith('.docx'):
                return Response({'error': 'Only DOCX files are supported'}, status=400)
            
            # Extract text using python-docx
            doc = docx.Document(material.file.path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"
            
            return Response({'text': text})
        
        except Exception as e:
            return Response({'error': str(e)}, status=500)


    @action(detail=False, methods=['post'], url_path='extract-text-from-url')
    def extract_text_from_url(self, request):
        """Extract text from a DOCX file URL."""
        file_url = request.data.get('file_url')
        material_id = request.data.get('material_id')
        
        if not file_url and not material_id:
            return Response({"error": "No file URL or material ID provided"}, 
                           status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Get the document directly by ID if provided
            if material_id:
                document = Document.objects.get(id=material_id)
            else:
                # Try to extract document ID from URL
                # This is a fallback method and might need adjustment
                try:
                    # Handle different URL formats
                    if '/' in file_url:
                        # Extract the last part of the URL path
                        path_parts = file_url.split('/')
                        filename = path_parts[-1]
                        # Remove file extension if present
                        if '.' in filename:
                            document_id = filename.split('.')[0]
                        else:
                            document_id = filename
                    else:
                        document_id = file_url
                        
                    # Try to convert to integer if it's numeric
                    try:
                        document_id = int(document_id)
                    except ValueError:
                        pass
                        
                    document = Document.objects.get(id=document_id)
                except Exception as e:
                    return Response({"error": f"Could not extract document ID from URL: {str(e)}"}, 
                                   status=status.HTTP_400_BAD_REQUEST)
            
            # Check if the user has access to this document
            if document.classroom.creator != request.user:
                return Response({"error": "You don't have permission to access this file"}, 
                              status=status.HTTP_403_FORBIDDEN)
            
            # Check if it's a DOCX file
            if not document.file.name.lower().endswith('.docx'):
                return Response({"error": "Only DOCX files are supported"}, 
                              status=status.HTTP_400_BAD_REQUEST)
            
            # Get the file path
            file_path = document.file.path
            
            # Extract text from DOCX
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
                
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            full_text.append(cell.text)
            
            return Response({"text": "\n".join(full_text)})
        
        except Document.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

